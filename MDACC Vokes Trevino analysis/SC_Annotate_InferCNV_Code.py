#!/usr/bin/env python
# coding: utf-8

#import packages

import numpy as np
import pandas as pd
import scanpy as sc
import matplotlib.pyplot as plt
import os


sc.settings.verbosity = 3             # verbosity: errors (0), warnings (1), info (2), hints (3)
sc.logging.print_header()
sc.settings.set_figure_params(dpi=80, facecolor='white')
pd.options.display.max_columns = None


#input/output directories

directory = "/Patient_Samples"
metadirectory = "/Le_Vokes_sample_collection/"


# # Read Data


adata = sc.read('/Vokes_Le_scRNA_Immune/2024_02_23_adata_afterQC.h5ad'

adata.shape


# # General UMAP and Annotation


adata.X = adata.raw.X


#Identify highly variable genes (no batch key)
sc.pp.highly_variable_genes(adata,max_mean=10, flavor = "seurat")


sc.pp.scale(adata, max_value=10)


# PCA
intialization = 3120
sc.pp.pca(adata, random_state=intialization, n_comps = 100 ,svd_solver='arpack')


sc.pp.neighbors(adata, random_state=intialization, n_neighbors=30, n_pcs=100,  method = 'umap', metric = 'euclidean')


#UMAP
sc.tl.umap(adata, random_state=intialization,spread = 1)


#individual markers
sc.pl.umap(adata, color=['EPCAM','KRT18','PTPRC','CD3E','CD19','CD33','BGN','CDH5'],title=["EPCAM (Epithelial)",'KRT18 (Epithelial)','PTCPRC (Lymphocyte)','CD3E (T cells)', 'CD19 (B cells)', 'CD33 (Myeloid)','BGN (Fibroblast)','CDH5 (Endothelial)'],show = True, size=20,color_map = 'viridis')


sc.tl.leiden(adata, random_state=intialization, resolution=0.1,use_weights = True, key_added = 'leiden_0.1')


sc.pl.umap(adata, color=['leiden_0.1'],show = True, size=10,legend_loc = 'on data')


#known signatures
mg_dot_dict = {}
mg_dot_dict['IMMUNE'] = ['PTPRC','CD3G','CD3E','CD79A','BLNK', 'CD68', 'CSF1R', 'MARCO','CD207']
mg_dot_dict['ENDOTHELIAL'] = ['PECAM1','CD34', 'VWF','CDH5','TIE1']
mg_dot_dict['FIBROBLAST'] =['ACTA2', 'MCAM', 'MYLK', 'MYL9', 'FAP', 'THY1']
mg_dot_dict['EPITHELIAL'] = ['EPCAM','SFN','KRT19','KRT18','KRT8','CD24']


sc.pl.dotplot(adata,mg_dot_dict,groupby='leiden_0.1',cmap='viridis',standard_scale='var',mean_only_expressed=False, dendrogram=True)


#manual assignment from dotplot and individual markers
OverlapDict = {}
OverlapDict['0'] = 'IMMUNE'#'LYMPHOID'
OverlapDict['1'] = 'IMMUNE'#'LYMPHOID'
OverlapDict['2'] = 'FIBROBLAST'
OverlapDict['3'] = 'IMMUNE'#'MYELOID'
OverlapDict['4'] = 'IMMUNE'#'LYMPHOID'
OverlapDict['5'] = 'EPITHELIAL'
OverlapDict['6'] = 'ENDOTHELIAL'
OverlapDict['7'] = 'IMMUNE'#'MYELOID'
OverlapDict['8'] = 'EPITHELIAL'
OverlapDict['9'] = 'IMMUNE'#'LYMPHOID'
OverlapDict['10'] = 'EPITHELIAL'
OverlapDict['11'] = 'IMMUNE'#'LYMPHOID'
OverlapDict['12'] = 'EPITHELIAL'
OverlapDict['13'] = 'EPITHELIAL'
OverlapDict['14'] = 'EPITHELIAL'
OverlapDict['15'] = 'EPITHELIAL'
OverlapDict['16'] = 'ENDOTHELIAL'
OverlapDict['17'] = 'IMMUNE'

adata.obs['FINAL_Cluster'] = adata.obs['leiden_0.1'].map(OverlapDict).astype('category')


# # INFERCNV outputs


#get patient numbers for adata subset with greater than two EPITHELIAL cells
adata_sub = adata[(adata.obs["FINAL_Cluster"].isin(['EPITHELIAL'])) , :].copy()
PatientNo = adata_sub.obs['PATIENT_NO'].value_counts().index.tolist()
Values = adata_sub.obs['PATIENT_NO'].value_counts().values.tolist()

threshold = 2
new = []
for i in range(0,len(Values)):
    if Values[i] > threshold:
        new.append(PatientNo[i])
PatientNo=new


#remove ATAC patient no
todel=[patient_08]
for p in todel:
    del PatientNo[PatientNo.index(p)]


samples=PatientNo


#make nonATAC outputs: use endothelial and fibroblast as reference
for i in range(0, len(samples)):
    print(samples[i])
    slist=[]
    slist.append(samples[i])
    #adata_sub = adata[ ((adata.obs["FINAL_Cluster"].isin(['MYELOID']))  | (  ~(adata.obs['SEQBATCH'].isin(['ATAC_batch_2022_10']))  &  (adata.obs['PATIENT_NO'].isin(slist)) &  (adata.obs["FINAL_Cluster"].isin(['EPITHELIAL'])) )), :].copy()
    adata_sub = adata[ ((adata.obs["FINAL_Cluster"].isin(['ENDOTHELIAL','FIBROBLAST']) & ~(adata.obs['SEQBATCH'].isin(['ATAC_batch_2022_10']))) | ( (adata.obs['PATIENT_NO'].isin(slist)) &  (adata.obs["FINAL_Cluster"].isin(['EPITHELIAL']))  & ~(adata.obs['SEQBATCH'].isin(['ATAC_batch_2022_10']))    )), :].copy()
    #print(adata_sub.obs['SEQBATCH'].value_counts())
    #adata_sub = adata[ (adata.obs["FINAL_Cluster"].isin(['NK','T CELL','MALIGNANT'])) & (adata.obs['SampleLegend'].isin(slist)), :].copy()
    #Annotation_df = pd.DataFrame(data = adata_sub.obs[["LYMPHOID_sub",'PATIENT_NO']], index = adata_sub.obs.index.tolist())
    Annotation_df = pd.DataFrame(data = adata_sub.obs[["FINAL_Cluster",'BatchLegend']], index = adata_sub.obs.index.tolist())
    Annotation_df[''] = Annotation_df.apply (lambda row: label_annotate(row), axis=1)
    del Annotation_df["FINAL_Cluster"]
    #del Annotation_df['PATIENT_NO']
    del Annotation_df['BatchLegend']
    Annotation_df.to_csv('/Vokes_Le_scRNA_Immune/INFERCNV_R_PATIENT/annotation_'+str(samples[i])+'_NonATAC.txt', sep="\t", header=False)
    INFERCNV_count_df = adata_sub.to_df(layer="counts")
    INFERCNV_count_df = INFERCNV_count_df.transpose()
    INFERCNV_count_df.to_csv('/Vokes_Le_scRNA_Immune/INFERCNV_R_PATIENT/INFERCNV_input_'+str(samples[i])+'_NonATAC.csv', sep="\t")

#ATAC samples
PatientNo=[patient_08]
samples=PatientNo


#make ATAC sample outputs: use endothelial and fibroblast as reference
for i in range(0, len(samples)):
    print(samples[i])
    slist=[]
    slist.append(samples[i])
    adata_sub = adata[ ((adata.obs["FINAL_Cluster"].isin(['ENDOTHELIAL','FIBROBLAST']) & ~(adata.obs['SEQBATCH'].isin(['ATAC_batch_2022_10']))) | ( (adata.obs['PATIENT_NO'].isin(slist)) &  (adata.obs["FINAL_Cluster"].isin(['EPITHELIAL']))  & (adata.obs['SEQBATCH'].isin(['ATAC_batch_2022_10']))    )), :].copy()
    Annotation_df = pd.DataFrame(data = adata_sub.obs[["FINAL_Cluster",'BatchLegend']], index = adata_sub.obs.index.tolist())
    Annotation_df[''] = Annotation_df.apply (lambda row: label_annotate(row), axis=1)
    del Annotation_df["FINAL_Cluster"]
    del Annotation_df['BatchLegend']
    Annotation_df.to_csv('/Vokes_Le_scRNA_Immune/INFERCNV_R_PATIENT/annotation_'+str(samples[i])+'_ATAC.txt', sep="\t", header=False)
    INFERCNV_count_df = adata_sub.to_df(layer="counts")
    INFERCNV_count_df = INFERCNV_count_df.transpose()
    INFERCNV_count_df.to_csv('/Users/sgtrevino1/Documents/Vokes_Le_scRNA_Immune/INFERCNV_R_PATIENT/INFERCNV_input_'+str(samples[i])+'_ATAC.csv', sep="\t")


# # INFERCNV Results

#make epithelial UMAP of InferCNV cluster numbers
adata_sub = adata[(adata.obs["FINAL_Cluster"].isin(['EPITHELIAL'])) , :].copy()
adata_sub.X = adata_sub.raw.X
sc.pp.highly_variable_genes(adata_sub, min_mean=0.0125, max_mean=7, min_disp=0.5)
sc.pp.scale(adata_sub, max_value=10)


#get patient numbers for adata subset with greater than two EPITHELIAL cells
pat_counts = adata_sub.obs['PATIENT_NO'].value_counts()
keep = pat_counts.index[pat_counts > 2]
adata_sub = adata_sub[adata_sub.obs['PATIENT_NO'].isin(keep)]


#compute UMAP
sc.tl.pca(adata_sub, svd_solver='arpack')
init = 100
sc.pp.neighbors(adata_sub, random_state=init, n_neighbors=30, n_pcs=50,  method = 'umap',metric = 'euclidean')
sc.tl.umap(adata_sub, random_state=init,spread = 2)


#get sample numbers as string
samples = adata_sub.obs['PATIENT_NO'].value_counts().index.tolist()
#add atac run
samples.append('502088_ATAC')
#make string
samplestr = []
for s in samples:
    samplestr.append(str(s))
new = []
for s in samplestr:
    new.append(s.split('.')[0])
samplestr = new


#check for infercnv results
import os.path
tsamples = samplestr
for i in range(0,len(tsamples)):
    filename = '/Vokes_Le_scRNA_Immune/INFERCNV_R_PATIENT_OUT_SCALE/InferCNV_'+str(tsamples[i])+'_out/infercnv.observation_groupings.txt'
    #print(filename)
    if not os.path.isfile(filename):
        print('Not Found')
        print(i+1)
        print(tsamples[i])

#create and save UMAP with InferCNV clusters for each sample
tsamples = samplestr
for i in range(0,len(tsamples)):
    filename = '/Vokes_Le_scRNA_IMMUNE/INFERCNV_R_PATIENT_OUT_SCALE/InferCNV_'+str(tsamples[i])+'_out/infercnv.observation_groupings.txt'
    
    if os.path.isfile(filename):
        assign = pd.read_csv(filename,sep=' ')
        
        infercnvr = [np.nan]*len(adata_sub.obs.index.tolist())
        cellsize = [10]*len(adata_sub.obs.index.tolist())
        colorcount = assign.groupby(['Dendrogram Group','Dendrogram Color'],as_index=False).size()
        colordict={}
        color = colorcount['Dendrogram Color'].tolist()
        groupnum = colorcount['Dendrogram Group'].astype(str).tolist()
        for j in range(0,len(color)):
            colordict[groupnum[j]] = color[j]
            
        dataindex = adata_sub.obs.index.tolist()
        
        cells = assign.index.tolist()
        
        ctype = assign['Dendrogram Group'].astype(str).tolist()
        for j in range(0, len(cells)):
            if cells[j].replace('.','-') in dataindex:
                infercnvr[dataindex.index(cells[j].replace('.','-'))] = ctype[j]
                cellsize[dataindex.index(cells[j].replace('.','-'))] = 15
        adata_sub.obs["InfercnvR"] = infercnvr
        plt.rcParams['figure.figsize']=(4,4)
        sc.pl.umap(adata_sub, color=['InfercnvR'],show = True,add_outline=False,outline_width=(0.05,0.9),na_color='lightgrey', palette=colordict,title = "",size=cellsize,frameon=False, legend_fontsize=10, legend_fontoutline=2, save = 'InferCNV_'+str(tsamples[i])+'.png')                                                                  
        
    
    else:
        print('Not Found')
        print(i+1)
        print(samples[i])


# # InferCNV initial Tumor/Normal Assignment


geneorder_df  = pd.read_csv('/Vokes_Le_scRNAseq_2020cohort/INFERCNV_R/geneorder.txt',header=None,sep='\t')



#fill in cnv_df with InferCNV copy number variation
tsamples = samplestr
li = []
maxlen = 0
for i in range(0,len(tsamples)):
#for i in range(0,1):
    #print(tsamples[i])
    filename = '/Vokes_Le_scRNA_IMMUNE/INFERCNV_R_PATIENT_OUT_SCALE/InferCNV_'+str(tsamples[i])+'_out/infercnv.observations.txt'
    if os.path.isfile(filename):
        temp_df = pd.read_csv(filename,sep=' ')
        temp_df = temp_df.transpose()
        
        if temp_df.isnull().values.any():
            print(f"nan in {tsamples[i]}")
        #print(temp_df.shape)
        li.append(temp_df)
cnv_df = pd.concat(li, axis=0, ignore_index=False)



#fill in missing cnv with variance of 1.0
cnv_df.fillna(1.0,inplace=True)


#calculate a UMAP based on the copy number variation
cnvdata = sc.AnnData(cnv_df)
# Set seed
intialization = 3120
sc.pp.pca(cnvdata, random_state=intialization, n_comps = 100 ,svd_solver='arpack')
sc.pl.pca_variance_ratio(cnvdata, n_pcs= 100, log=True, show = True)
sc.pl.pca(cnvdata)

sc.pp.neighbors(cnvdata, random_state=intialization, n_neighbors=15, n_pcs=100,  method = 'umap', metric = 'euclidean')
sc.tl.umap(cnvdata, random_state=intialization, spread=1.0)


#cluster the data 
sc.tl.leiden(cnvdata, random_state=intialization, resolution=0.3,use_weights = True)
sc.pl.umap(cnvdata, color=['leiden'],legend_loc= 'on data',frameon=False, title="")


#create a heatmap of cnv for each cluster
sc.tl.dendrogram(cnvdata,groupby='leiden')
sc.pl.heatmap(cnvdata,var_names=genes,groupby=['leiden'],var_group_labels=['cnv_status'],cmap='coolwarm',dendrogram=True,figsize=(10,25), save = 'InferCNVR_heatmap')



#based on the heatmap cluster 0 has the lowest cnv and is intially labeled as normal 
cnvdata.obs["cnv_status"] = "tumor"
cnvdata.obs.loc[
    cnvdata.obs["leiden"].isin(['0']), "cnv_status"
] = "normal"



#calculate the cnv_score of each cell 
avgcol = []
for i in range(0,len(cnv_df.index.tolist())):
    current_row = cnv_df.iloc[i].to_list()
    avg_row=[]
    for j in range(0,len(current_row)):
        if current_row[j] != 1.0:
            avg_row.append(abs(current_row[j]-1.0))
    avgcol.append(sum(avg_row)/len(avg_row))



cnvdata.obs['cnv_score'] = avgcol



#calculate the cnv_avg of each cluster
numgroups = max(cnvdata.obs['leiden'].value_counts().index.astype(int).tolist()) +1
avg_dict = {}
for i in range(0,numgroups):
    avg_dict[str(i)] = sum(cnvdata.obs[cnvdata.obs['leiden'] == str(i)]['cnv_score'].tolist())/(len(cnvdata.obs[cnvdata.obs['leiden'] == str(i)]['cnv_score'].tolist()))



cnvdata.obs['cnv_avg'] = cnvdata.obs['leiden'].map(avg_dict).astype('float')



sc.pl.umap(cnvdata, color=['cnv_score','cnv_avg','cnv_status'],cmap='viridis')


#transfer normal/tumor labels to main adata
infercnvr = [np.nan]*len(adata.obs.index.tolist())
dataindex = adata.obs.index.tolist()
cells = cnvdata.obs.index.tolist()
ctype = cnvdata.obs['cnv_status'].astype(str).tolist()
for j in range(0, len(cells)):
    if cells[j].replace('.','-') in dataindex:
        infercnvr[dataindex.index(cells[j].replace('.','-'))] = ctype[j]
adata.obs["InfercnvR_status"] = infercnvr



#get epithelial subset and make UMAP of normal/tumor labels for each sample
#get adata subset
adata_sub = adata[(adata.obs["FINAL_Cluster"].isin(['EPITHELIAL'])) , :].copy()
adata_sub.X = adata_sub.raw.X
sc.pp.highly_variable_genes(adata_sub, min_mean=0.0125, max_mean=7, min_disp=0.5)
sc.pp.scale(adata_sub, max_value=10)

pat_counts = adata_sub.obs['PATIENT_NO'].value_counts()
keep = pat_counts.index[pat_counts > 2]
adata_sub = adata_sub[adata_sub.obs['PATIENT_NO'].isin(keep)]

sc.tl.pca(adata_sub, svd_solver='arpack')
init = 100
sc.pp.neighbors(adata_sub, random_state=init, n_neighbors=30, n_pcs=50,  method = 'umap',metric = 'euclidean')
sc.tl.umap(adata_sub, random_state=init,spread = 2)


#create and save UMAP with normal/tumor annotations for each sample
tsamples = samplestr
for i in range(0,len(tsamples)):
    filename = '/Users/sgtrevino1/Documents/Vokes_Le_scRNA_IMMUNE/INFERCNV_R_PATIENT_OUT_SCALE/InferCNV_'+str(tsamples[i])+'_out/infercnv.observation_groupings.txt'
    print(filename)
    if os.path.isfile(filename):
        assign = pd.read_csv(filename,sep=' ')
        
        infercnvr = [np.nan]*len(adata_sub.obs.index.tolist())
        infercnvr_results = adata_sub.obs["InfercnvR_status"].tolist()
        cellsize = [10]*len(adata_sub.obs.index.tolist())
        colorcount = assign.groupby(['Dendrogram Group','Dendrogram Color'],as_index=False).size()
        colordict={}
        color = colorcount['Dendrogram Color'].tolist()
        groupnum = colorcount['Dendrogram Group'].astype(str).tolist()
        for j in range(0,len(color)):
            colordict[groupnum[j]] = color[j]
            
        dataindex = adata_sub.obs.index.tolist()
        
        cells = assign.index.tolist()
        
        ctype = assign['Dendrogram Group'].astype(str).tolist()
        for j in range(0, len(cells)):
            if cells[j].replace('.','-') in dataindex:
                infercnvr[dataindex.index(cells[j].replace('.','-'))] = infercnvr_results[dataindex.index(cells[j].replace('.','-'))]#ctype[j]
                cellsize[dataindex.index(cells[j].replace('.','-'))] = 15
        adata_sub.obs["InfercnvR"] = infercnvr
        plt.rcParams['figure.figsize']=(4,4)
        sc.pl.umap(adata_sub, color=['InfercnvR'],show = True,add_outline=False,outline_width=(0.05,0.9),na_color='lightgrey',title = "",size=cellsize,frameon=False, legend_fontsize=10, legend_fontoutline=2, save = 'InferCNV_'+str(tsamples[i])+'_TumorNormal.png')                                                                  
       
        print('Not Found')
        print(i+1)
        print(tsamples[i])


# # InferCNV manual inspection and correction


#combine all inferCNV results and annotations into docx file for manual inspection/correction
samples = adata_sub.obs['PATIENT_NO'].value_counts().index.tolist()
new=[]
for sample in samples:
    new.append(int(sample))
samples=new

#make documents
from docx import Document
from docx.shared import Inches
import os
document = Document()
for i in range(0,len(samples)):
    
    Batches = adata[(adata.obs["FINAL_Cluster"].isin(['EPITHELIAL']))&(adata.obs["PATIENT_NO"].isin([float(samples[i])])),:].obs['BatchLegend'].value_counts().index.tolist()
    print(Batches)
    MergeBatches = [', '.join(Batches[0 : len(Batches)])]    
    Heading = f"Patient {str(samples[i])}\t Batch:{MergeBatches[0]}"
    Image1path = '/Vokes_Le_scRNA_IMMUNE/INFERCNV_R_PATIENT_OUT_SCALE/InferCNV_' +str(samples[i])+'_out/infercnv.png'
    Image2path = '/Vokes_Le_scRNA_IMMUNE/figures/umapinferCNV_' +str(samples[i])+'.png'
    Image3path = '/Vokes_Le_scRNA_IMMUNE/figures/umapinferCNV_' +str(samples[i])+'_TumorNormal.png'
   
    text = f"['normal',’normal’,’normal’,'normal'] #{samples[i]}"
    document.add_heading(Heading,level=2)
    if os.path.isfile(Image1path):
        document.add_picture(Image1path, width=Inches(6.0))
        #document.add_paragraph(text)
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(Image2path, width=Inches(3.0))
        if os.path.isfile(Image3path):
            run.add_picture(Image3path, width=Inches(3.0))
        
    document.add_page_break()
document.save('INFERCNV_R_PATIENT_SCALED_Results_IMMUNE.docx')



#correct values from manual inspection (patientno, groups, new assignment)
correct= [(patient_41,['3','4'],'tumor'),(patient_48,['1'],'normal')]
Icnv = adata.obs["InfercnvR_status"].tolist()
Pno = adata.obs["PATIENT_NO"].tolist()
Cluster = adata.obs['infercnv_cluster'].tolist()

for i in range(0,len(Icnv)):
    for j in range(0,len(correct)):
        if (Pno[i] == correct[j][0]) and (Cluster[i] in correct[j][1]):
            #print('found')
            Icnv[i] = correct[j][2]

adata.obs['infercnv_cluster_m'] = Icnv


# # Output Data


# Subset Data for analysis 


#take Epithelial malignant subset
adata = adata[ (adata.obs['infercnv_cluster_m'].isin(['tumor'])), :].copy()

keep = ['Biopsy-1','JH064','JH067','JH128','JH139','Lung-Tumor-10','JH-297','JH143','Lung-Tumor-1','Lung-Tumor-7','Lung-Tumor-8','NSTAR1-TumorA','JH-304','JH-305','JH033','JH038','JH104','JH095']

adata = adata[adata.obs['Specimen_ID'].isin(keep)]

#save data
adata.write('Vokes_Le_scRNA_EGFR_5_25/CodeForBen/EGFRSub.h5ad')

